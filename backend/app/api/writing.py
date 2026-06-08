"""写作辅助 API。"""

import logging
from typing import List, Optional
from fastapi import APIRouter, HTTPException, Depends, Query
from pydantic import BaseModel, Field
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.session import get_db
from app.services.writing_service import WritingAssistantService
from app.core.security import get_optional_user

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/writing", tags=["写作"])


class RecommendRequest(BaseModel):
    text: str = Field(..., description="需要引用的文本段落")
    top_k: int = Field(default=5, ge=1, le=10)


class CitationResult(BaseModel):
    paper_id: Optional[str] = None
    title: str
    authors: str
    year: Optional[int]
    arxiv_id: Optional[str]
    doi: Optional[str]
    abstract_snippet: str
    similarity: float
    bibtex: str
    role: Optional[str] = None
    role_label: Optional[str] = None
    role_reason: Optional[str] = None
    match_score: Optional[float] = None
    match_status: Optional[str] = None
    match_label: Optional[str] = None
    match_terms: Optional[List[str]] = None
    match_explanation: Optional[str] = None
    decision_label: Optional[str] = None
    decision_action: Optional[str] = None
    decision_warning: Optional[str] = None
    decision_confidence: Optional[str] = None


class RelatedWorkRequest(BaseModel):
    topic: str = Field(..., description="研究主题")
    max_papers: int = Field(default=5, ge=2, le=10)
    language: str = Field(default="chinese")


class PolishRequest(BaseModel):
    text: str = Field(..., description="需要润色的文本")
    style: str = Field(default="academic", description="润色风格: academic, concise, fluent, english")


class AbstractRequest(BaseModel):
    title: str = Field(..., description="论文标题")
    key_points: str = Field(..., description="关键要点")
    language: str = Field(default="chinese")


class TextResponse(BaseModel):
    result: str


@router.post("/recommend-citations", response_model=List[CitationResult])
async def recommend_citations(req: RecommendRequest, db: AsyncSession = Depends(get_db)):
    """根据写作内容推荐引用论文。"""
    service = WritingAssistantService(db)
    return await service.recommend_citations(req.text, top_k=req.top_k)


@router.post("/related-work", response_model=TextResponse)
async def generate_related_work(req: RelatedWorkRequest, db: AsyncSession = Depends(get_db)):
    """基于知识库生成 Related Work 章节。"""
    service = WritingAssistantService(db)
    result = await service.generate_related_work(
        req.topic, max_papers=req.max_papers, language=req.language
    )
    return TextResponse(result=result)


@router.post("/related-work/table")
async def generate_related_work_table(req: RelatedWorkRequest, db: AsyncSession = Depends(get_db)):
    """生成 Related Work 对比表。"""
    service = WritingAssistantService(db)
    return await service.generate_related_work_table(req.topic, max_papers=req.max_papers)


@router.post("/polish", response_model=TextResponse)
async def polish_text(req: PolishRequest, db: AsyncSession = Depends(get_db)):
    """润色学术文本。"""
    service = WritingAssistantService(db)
    result = await service.polish_text(req.text, style=req.style)
    return TextResponse(result=result)


@router.post("/generate-abstract", response_model=TextResponse)
async def generate_abstract(req: AbstractRequest, db: AsyncSession = Depends(get_db)):
    """生成论文摘要。"""
    service = WritingAssistantService(db)
    result = await service.generate_abstract(req.title, req.key_points, req.language)
    return TextResponse(result=result)


class LitReviewRequest(BaseModel):
    topic: str = Field(..., description="研究方向")
    max_papers: int = Field(default=10, ge=3, le=20)
    language: str = Field(default="chinese")


class CompareRequest(BaseModel):
    paper_ids: List[str] = Field(..., min_length=2, max_length=5, description="要对比的论文 ID 列表")


@router.post("/literature-review")
async def generate_literature_review(req: LitReviewRequest, db: AsyncSession = Depends(get_db)):
    """生成完整文献综述。"""
    service = WritingAssistantService(db)
    result = await service.generate_literature_review(req.topic, req.language, req.max_papers)
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])
    return result


@router.post("/compare-papers")
async def compare_papers(req: CompareRequest, db: AsyncSession = Depends(get_db)):
    """对比分析多篇论文。"""
    service = WritingAssistantService(db)
    result = await service.compare_papers(req.paper_ids)
    return TextResponse(result=result)


@router.post("/export-bibtex")
async def export_bibtex(paper_ids: List[str], db: AsyncSession = Depends(get_db)):
    """导出多篇论文的 BibTeX。"""
    service = WritingAssistantService(db)
    bibtex = await service.export_bibtex(paper_ids)
    return {"bibtex": bibtex}


class ExportRequest(BaseModel):
    format: str = Field(default="bibtex", description="导出格式: bibtex, csv, markdown")
    paper_ids: Optional[List[str]] = None
    project_id: Optional[str] = None


class ReportRequest(BaseModel):
    paper_ids: List[str] = Field(..., min_length=1, description="论文 ID 列表")
    title: str = Field(default="组会报告", description="报告主标题")
    custom_prompt: Optional[str] = Field(default=None, max_length=4000, description="自定义组会报告生成要求")


class FeishuReportRequest(BaseModel):
    paper_ids: List[str] = Field(..., min_length=1)
    title: str = Field(default="组会报告")
    feishu_url: str = Field(..., description="飞书文档链接")
    custom_prompt: Optional[str] = Field(default=None, max_length=4000, description="自定义组会报告生成要求")


REPORT_EAST_ASIA_FONT = "宋体"
REPORT_LATIN_FONT = "Times New Roman"


def _set_run_fonts(run, east_asia: str = REPORT_EAST_ASIA_FONT, latin: str = REPORT_LATIN_FONT) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def _set_style_fonts(style, east_asia: str = REPORT_EAST_ASIA_FONT, latin: str = REPORT_LATIN_FONT) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    style.font.name = latin
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def _apply_paragraph_fonts(paragraph) -> None:
    for run in paragraph.runs:
        _set_run_fonts(run)


def _configure_report_doc_fonts(doc) -> None:
    from docx.shared import Pt

    _set_style_fonts(doc.styles["Normal"])
    doc.styles["Normal"].font.size = Pt(11)


def _add_report_heading(doc, text: str, level: int):
    paragraph = doc.add_heading(text, level=level)
    _apply_paragraph_fonts(paragraph)
    return paragraph


def _add_report_paragraph(doc, text: str = "", style: str | None = None):
    paragraph = doc.add_paragraph(text, style=style)
    _apply_paragraph_fonts(paragraph)
    return paragraph


def _add_markdown_report(doc, markdown: str) -> None:
    for raw_line in (markdown or "").splitlines():
        line = raw_line.strip()
        if not line:
            _add_report_paragraph(doc, "")
        elif line.startswith("### "):
            _add_report_heading(doc, line[4:].strip(), level=3)
        elif line.startswith("## "):
            _add_report_heading(doc, line[3:].strip(), level=2)
        elif line.startswith("# "):
            _add_report_heading(doc, line[2:].strip(), level=1)
        elif line.startswith("- "):
            _add_report_paragraph(doc, line[2:].strip(), style="List Bullet")
        else:
            _add_report_paragraph(doc, line)


@router.post("/group-report")
async def generate_group_report(req: ReportRequest, db: AsyncSession = Depends(get_db)):
    """生成组会报告（结构化总结 + Word 下载）。"""
    from app.services.report_service import ReportService
    service = ReportService(db)
    result = await service.generate_report(req.paper_ids, req.title, custom_prompt=req.custom_prompt)
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])

    # 生成 Word 文档
    from docx import Document
    import io
    from datetime import datetime

    doc = Document()
    _configure_report_doc_fonts(doc)

    _add_report_heading(doc, req.title, level=0)
    _add_report_paragraph(doc, f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    _add_report_paragraph(doc, f"共 {result['paper_count']} 篇论文")
    _add_report_paragraph(doc, "")

    if result.get("custom_report"):
        _add_markdown_report(doc, result["custom_report"])
    else:
        for i, pdata in enumerate(result["papers"], 1):
            _add_report_heading(doc, f"{pdata['title']}", level=1)
            _add_report_paragraph(doc, f"作者: {pdata['authors']}  |  年份: {pdata['year']}  |  arXiv: {pdata.get('arxiv_id', 'N/A')}")

            for section_name, section_content in pdata["sections"].items():
                if section_content:
                    _add_report_heading(doc, section_name, level=2)
                    # 解析 Markdown 要点为 Word 列表
                    for line in section_content.split("\n"):
                        line = line.strip()
                        if line.startswith("- "):
                            _add_report_paragraph(doc, line[2:], style='List Bullet')
                        elif line:
                            _add_report_paragraph(doc, line)

            _add_report_paragraph(doc, "")

    # 输出为 bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    from fastapi.responses import StreamingResponse
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=group_report_{datetime.now().strftime('%Y%m%d')}.docx"},
    )


@router.get("/group-report-md")
async def generate_report_markdown(
    paper_ids: str = Query(..., description="论文 ID，逗号分隔"),
    title: str = Query(default="组会报告"),
    custom_prompt: Optional[str] = Query(default=None, max_length=4000),
    db: AsyncSession = Depends(get_db),
):
    """生成组会报告 Markdown 文本（可直接粘贴到飞书/Notion）。"""
    ids = [p.strip() for p in paper_ids.split(",") if p.strip()]
    from app.services.report_service import ReportService
    service = ReportService(db)
    result = await service.generate_report(ids, title, custom_prompt=custom_prompt)
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])

    md = f"# {result['title']}\n\n生成时间: {result['generated_at'][:19]}  |  共 {result['paper_count']} 篇\n\n---\n\n"
    if result.get("custom_report"):
        md += result["custom_report"].strip() + "\n"
    else:
        for i, pdata in enumerate(result["papers"], 1):
            md += f"## {i}. {pdata['title']}\n\n"
            md += f"**作者**: {pdata['authors']}  |  **年份**: {pdata['year']}  |  **arXiv**: {pdata.get('arxiv_id', 'N/A')}\n\n"
            for section_name, section_content in pdata["sections"].items():
                if section_content:
                    md += f"### {section_name}\n\n{section_content}\n\n"
            md += "---\n\n"

    return TextResponse(result=md)


@router.post("/group-report-to-feishu")
async def write_report_to_feishu(req: FeishuReportRequest, db: AsyncSession = Depends(get_db)):
    """生成组会报告并写入飞书文档。"""

    # 1. 生成报告
    ids = req.paper_ids
    from app.services.report_service import ReportService
    service = ReportService(db)
    result = await service.generate_report(ids, req.title, custom_prompt=req.custom_prompt)
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])

    # 2. 生成 Markdown
    md = f"# {result['title']}\n\n生成时间: {result['generated_at'][:19]}  |  共 {result['paper_count']} 篇\n\n---\n\n"
    if result.get("custom_report"):
        md += result["custom_report"].strip() + "\n"
    else:
        for i, pdata in enumerate(result["papers"], 1):
            md += f"## {i}. {pdata['title']}\n\n"
            md += f"**作者**: {pdata['authors']}  |  **年份**: {pdata['year']}  |  **arXiv**: {pdata.get('arxiv_id', 'N/A')}\n\n"
            for section_name, section_content in pdata["sections"].items():
                if section_content:
                    md += f"### {section_name}\n\n{section_content}\n\n"
            md += "---\n\n"

    # 3. 写入飞书
    try:
        from app.services.feishu_service import FeishuService
        feishu = FeishuService()
        if not feishu.configured:
            return {"status": "fallback", "message": "飞书未配置，请设置 FEISHU_APP_ID 和 FEISHU_APP_SECRET", "markdown": md}

        doc_id = FeishuService.parse_doc_id(req.feishu_url)
        await feishu.append_blocks(doc_id, md)
        await feishu.close()
        return {"status": "success", "message": "报告已写入飞书文档", "doc_id": doc_id}
    except ValueError as e:
        return {"status": "error", "message": str(e), "markdown": md}
    except Exception as e:
        logger.error(f"飞书写入失败: {e}")
        return {"status": "error", "message": f"飞书写入失败: {str(e)}", "markdown": md}


@router.get("/group-report-json")
async def generate_report_json(
    paper_ids: str = Query(..., description="论文 ID，逗号分隔"),
    title: str = Query(default="组会报告"),
    custom_prompt: Optional[str] = Query(default=None, max_length=4000),
    db: AsyncSession = Depends(get_db),
):
    """生成组会报告 JSON（用于前端预览）。"""
    ids = [p.strip() for p in paper_ids.split(",") if p.strip()]
    from app.services.report_service import ReportService
    service = ReportService(db)
    result = await service.generate_report(ids, title, custom_prompt=custom_prompt)
    return result


@router.post("/export")
async def export_data(req: ExportRequest, db: AsyncSession = Depends(get_db), current_user=Depends(get_optional_user)):
    """导出论文或研究数据。"""
    from uuid import UUID
    from sqlalchemy import select
    from app.db.models.paper import Paper
    from app.db.models.research import ResearchProject, ResearchIdea
    import csv, io

    service = WritingAssistantService(db)

    if req.format == "bibtex":
        if req.paper_ids:
            bibtex = await service.export_bibtex(req.paper_ids)
            return {"data": bibtex, "format": "bibtex"}
        else:
            result = await db.execute(select(Paper).limit(50))
            papers = result.scalars().all()
            entries = [service._generate_bibtex(p) for p in papers]
            return {"data": "\n\n".join(entries), "format": "bibtex"}

    elif req.format == "csv":
        result = await db.execute(select(Paper))
        papers = result.scalars().all()
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(["Title", "Authors", "Year", "Abstract", "arXiv ID", "DOI", "Tags"])
        for p in papers:
            writer.writerow([
                p.title,
                ", ".join(p.authors) if isinstance(p.authors, list) else str(p.authors),
                p.year, p.abstract[:200] if p.abstract else "",
                p.arxiv_id, p.doi,
                ", ".join(p.tags) if p.tags else "",
            ])
        return {"data": output.getvalue(), "format": "csv"}

    elif req.format == "markdown":
        ideas_text = ""
        if req.project_id:
            if not current_user:
                raise HTTPException(status_code=401, detail="请先登录")
            try:
                project_id = UUID(req.project_id)
            except ValueError:
                raise HTTPException(status_code=400, detail="Invalid project_id")
            project = (await db.execute(
                select(ResearchProject).where(
                    ResearchProject.id == project_id,
                    ResearchProject.user_id == current_user.id,
                )
            )).scalar_one_or_none()
            if not project:
                raise HTTPException(status_code=404, detail="项目未找到")
            result = await db.execute(
                select(ResearchIdea).where(ResearchIdea.project_id == project_id)
            )
            ideas = result.scalars().all()
            ideas_text = "\n\n".join([
                f"### {i.title}\n\n**可行性**: {i.feasibility_score}/10 | **创新性**: {i.novelty_score}/10\n\n{i.description or ''}\n\n"
                for i in ideas
            ])

        result = await db.execute(select(Paper))
        papers = result.scalars().all()
        papers_text = "\n".join([
            f"- [{p.title}](https://arxiv.org/abs/{p.arxiv_id}) ({p.year})"
            for p in papers if p.arxiv_id
        ])

        md = f"# 研究报告\n\n## 论文列表\n\n{papers_text}\n\n## 研究 Idea\n\n{ideas_text}"
        return {"data": md, "format": "markdown"}

    raise HTTPException(status_code=400, detail=f"Unsupported format: {req.format}")


# --- 申请书助手 ---

class GrantWriteRequest(BaseModel):
    section: str = Field(..., description="章节名称: 立项依据/研究内容/研究方案/特色创新/预期成果/研究基础")
    topic: str = Field(..., description="项目主题")
    background: str = Field(default="", description="项目背景/摘要")
    previous_content: str = Field(default="", description="前文已写的内容")


class GrantReviewRequest(BaseModel):
    section: str = Field(..., description="章节名称")
    content: str = Field(..., description="需评审的内容")
    topic: str = Field(default="")


class GrantInnovationRequest(BaseModel):
    topic: str = Field(...)
    background: str = Field(default="")
    methods: str = Field(default="")


class GrantPolishRequest(BaseModel):
    text: str = Field(..., description="需润色的文本")


@router.post("/grant/write-section", response_model=TextResponse)
async def grant_write_section(req: GrantWriteRequest, db: AsyncSession = Depends(get_db)):
    """撰写申请书章节（参考 eseckel/ai-for-grant-writing）。"""
    from app.services.grant_service import GrantService
    service = GrantService(db)
    result = await service.write_section(req.section, req.topic, req.background, req.previous_content)
    return TextResponse(result=result)


@router.post("/grant/review-section", response_model=TextResponse)
async def grant_review_section(req: GrantReviewRequest, db: AsyncSession = Depends(get_db)):
    """模拟 NSFC 评审专家审阅（参考 NSFC 评审模拟 Prompt）。"""
    from app.services.grant_service import GrantService
    service = GrantService(db)
    result = await service.review_section(req.section, req.content, req.topic)
    return TextResponse(result=result)


@router.post("/grant/extract-innovation", response_model=TextResponse)
async def grant_extract_innovation(req: GrantInnovationRequest, db: AsyncSession = Depends(get_db)):
    """提炼核心创新点（参考 NSFC 创新点提炼模板）。"""
    from app.services.grant_service import GrantService
    service = GrantService(db)
    result = await service.extract_innovation_points(req.topic, req.background, req.methods)
    return TextResponse(result=result)


@router.post("/grant/polish", response_model=TextResponse)
async def grant_polish(req: GrantPolishRequest, db: AsyncSession = Depends(get_db)):
    """润色申请书文本（参考 gpt_academic + eseckel clarity prompt）。"""
    from app.services.grant_service import GrantService
    service = GrantService(db)
    result = await service.polish_grant_text(req.text)
    return TextResponse(result=result)
