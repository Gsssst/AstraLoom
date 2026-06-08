"""写作项目管理服务 — 项目 CRUD、章节管理、模板系统、多格式导出。"""

import logging
import re
from typing import Any, List, Optional
from uuid import UUID

from sqlalchemy import select, func, update, delete
from sqlalchemy.ext.asyncio import AsyncSession

logger = logging.getLogger(__name__)


# 预定义写作模板
TEMPLATES = {
    "blank": {
        "name": "空白模板",
        "description": "从零开始的空白论文",
        "sections": [],
    },
    "acl": {
        "name": "ACL",
        "description": "ACL/EMNLP/NAACL 会议论文模板",
        "sections": [
            {"title": "Abstract", "level": 1},
            {"title": "Introduction", "level": 1},
            {"title": "Related Work", "level": 1},
            {"title": "Method", "level": 1},
            {"title": "Experiments", "level": 1},
            {"title": "Conclusion", "level": 1},
        ],
    },
    "cvpr": {
        "name": "CVPR",
        "description": "CVPR/ICCV/ECCV 会议论文模板",
        "sections": [
            {"title": "Abstract", "level": 1},
            {"title": "Introduction", "level": 1},
            {"title": "Related Work", "level": 1},
            {"title": "Method", "level": 1},
            {"title": "Experiments", "level": 1},
            {"title": "Conclusion", "level": 1},
        ],
    },
    "neurips": {
        "name": "NeurIPS",
        "description": "NeurIPS/ICML/ICLR 会议论文模板",
        "sections": [
            {"title": "Abstract", "level": 1},
            {"title": "Introduction", "level": 1},
            {"title": "Related Work", "level": 1},
            {"title": "Preliminaries", "level": 1},
            {"title": "Method", "level": 1},
            {"title": "Experiments", "level": 1},
            {"title": "Conclusion", "level": 1},
        ],
    },
    "icml": {
        "name": "ICML",
        "sections": [
            {"title": "Abstract", "level": 1},
            {"title": "Introduction", "level": 1},
            {"title": "Background", "level": 1},
            {"title": "Proposed Method", "level": 1},
            {"title": "Theoretical Analysis", "level": 1},
            {"title": "Experiments", "level": 1},
            {"title": "Conclusion", "level": 1},
        ],
    },
    "nsfc": {
        "name": "NSFC 申请书",
        "description": "国家自然科学基金申请书模板",
        "sections": [
            {"title": "立项依据与研究内容", "level": 1},
            {"title": "研究目标", "level": 1},
            {"title": "研究内容", "level": 1},
            {"title": "拟采取的研究方案", "level": 1},
            {"title": "技术路线", "level": 1},
            {"title": "可行性分析", "level": 1},
            {"title": "特色与创新之处", "level": 1},
            {"title": "预期成果", "level": 1},
            {"title": "研究基础", "level": 1},
        ],
    },
    "survey": {
        "name": "综述草稿",
        "description": "从研究方向生成文献综述草稿",
        "sections": [
            {"title": "Abstract", "level": 1},
            {"title": "Introduction", "level": 1},
            {"title": "Related Work", "level": 1},
            {"title": "Related Work Comparison Table", "level": 1},
            {"title": "Research Gaps", "level": 1},
            {"title": "References", "level": 1},
        ],
    },
}


class WritingProjectService:
    """写作项目管理服务。"""

    def __init__(self, session: AsyncSession):
        self.session = session

    # --- 模板 ---

    @staticmethod
    def get_templates() -> list:
        """获取可用模板列表。"""
        return [
            {"key": k, "name": v["name"], "description": v.get("description", ""),
             "section_count": len(v["sections"])}
            for k, v in TEMPLATES.items()
        ]

    @staticmethod
    def get_template_sections(template_type: str) -> list:
        """获取模板的预设章节。"""
        template = TEMPLATES.get(template_type, TEMPLATES["blank"])
        return template["sections"]

    # --- 项目 CRUD ---

    async def create_project(self, user_id: str, title: str,
                              description: str = "",
                              template_type: str = "blank",
                              metadata_json: Optional[dict] = None) -> dict:
        """创建新写作项目。"""
        from app.db.models.writing import WritingProject, WritingSection

        project = WritingProject(
            user_id=user_id,
            title=title,
            description=description,
            template_type=template_type,
            metadata_json=metadata_json or {},
        )
        self.session.add(project)
        await self.session.flush()

        # 从模板创建预设章节
        template_sections = self.get_template_sections(template_type)
        for i, sec in enumerate(template_sections):
            section = WritingSection(
                project_id=str(project.id),
                title=sec["title"],
                order=i,
            )
            self.session.add(section)

        await self.session.commit()
        await self.session.refresh(project)

        return self._project_to_dict(project)

    async def create_project_from_context(
        self,
        user_id: str,
        title: str,
        description: str = "",
        template_type: str = "blank",
        writing_type: str = "paper",
        research_project_id: Optional[str] = None,
        collection_ids: Optional[list[str]] = None,
        target_venue: str = "",
        target_year: str = "",
    ) -> dict:
        """Create a writing project bound to research direction and paper collections."""
        context = await self._resolve_writing_context(
            user_id=user_id,
            writing_type=writing_type,
            research_project_id=research_project_id,
            collection_ids=collection_ids or [],
            target_venue=target_venue,
            target_year=target_year,
        )
        seed_paper_ids = context.get("paper_ids") or []
        metadata = {
            "source": "context_bound_writing_project",
            "writing_context": context,
            "recommended_paper_ids": seed_paper_ids,
            "recommended_arxiv_ids": [],
            "evidence_status": "sufficient" if seed_paper_ids else "insufficient",
        }
        if target_venue or target_year:
            metadata["submission_profile"] = {
                "venue": (target_venue or "").strip(),
                "year": (target_year or "").strip(),
                "template_status": "missing",
                "status_label": "未绑定官方模板",
                "warnings": ["已设置投稿目标，但尚未上传或绑定官方模板。"],
            }

        effective_description = description or context.get("description") or ""
        project = await self.create_project(
            user_id=user_id,
            title=title,
            description=effective_description,
            template_type=template_type,
            metadata_json=metadata,
        )
        await self._seed_context_sections(project["id"], user_id, context)
        updated = await self.get_project(project["id"], user_id)
        return updated or project

    async def _resolve_writing_context(
        self,
        *,
        user_id: str,
        writing_type: str,
        research_project_id: Optional[str],
        collection_ids: list[str],
        target_venue: str,
        target_year: str,
    ) -> dict:
        from app.db.models.paper import Folder, PaperFolderItem
        from app.db.models.research import ResearchProject

        context: dict[str, Any] = {
            "writing_type": writing_type or "paper",
            "target_venue": (target_venue or "").strip(),
            "target_year": (target_year or "").strip(),
            "research_project_id": None,
            "research_project_name": "",
            "description": "",
            "collection_ids": [],
            "collection_names": [],
            "collection_sources": [],
            "paper_ids": [],
        }
        paper_ids: list[str] = []

        if research_project_id:
            try:
                rid = UUID(research_project_id)
            except ValueError as exc:
                raise ValueError("Invalid research_project_id") from exc
            result = await self.session.execute(
                select(ResearchProject).where(ResearchProject.id == rid, ResearchProject.user_id == UUID(user_id))
            )
            project = result.scalar_one_or_none()
            if not project:
                raise ValueError("研究方向未找到")
            context["research_project_id"] = str(project.id)
            context["research_project_name"] = project.name
            context["description"] = project.description or ""
            for paper_id in project.paper_ids or []:
                if paper_id and str(paper_id) not in paper_ids:
                    paper_ids.append(str(paper_id))

        seen_collections = set()
        for raw_id in collection_ids or []:
            if not raw_id or raw_id in seen_collections:
                continue
            seen_collections.add(raw_id)
            try:
                folder_id = UUID(raw_id)
            except ValueError as exc:
                raise ValueError("Invalid collection_id") from exc
            folder = (await self.session.execute(
                select(Folder).where(Folder.id == folder_id, Folder.user_id == UUID(user_id))
            )).scalar_one_or_none()
            if not folder:
                raise ValueError("论文分类未找到")
            paper_result = await self.session.execute(
                select(PaperFolderItem.paper_id)
                .where(PaperFolderItem.folder_id == folder.id, PaperFolderItem.user_id == UUID(user_id))
                .order_by(PaperFolderItem.created_at.desc())
            )
            folder_paper_ids = [str(pid) for pid in paper_result.scalars().all()]
            context["collection_ids"].append(str(folder.id))
            context["collection_names"].append(folder.name)
            context["collection_sources"].append({
                "id": str(folder.id),
                "name": folder.name,
                "paper_count": len(folder_paper_ids),
            })
            for paper_id in folder_paper_ids:
                if paper_id not in paper_ids:
                    paper_ids.append(paper_id)

        context["paper_ids"] = paper_ids
        return context

    async def _seed_context_sections(self, project_id: str, user_id: str, context: dict) -> None:
        project = await self.get_project(project_id, user_id)
        if not project:
            return
        sections_by_title = {section["title"].lower(): section for section in project.get("sections", [])}
        content_by_title = await self._build_context_seed_sections(project, user_id, context)
        for title, content in content_by_title.items():
            section = sections_by_title.get(title.lower())
            if section and not (section.get("content") or "").strip():
                await self.update_section(section["id"], user_id, content=content, status="draft")

    async def _build_context_seed_sections(self, project: dict, user_id: str, context: dict) -> dict[str, str]:
        evidence_table = await self.build_evidence_related_work_table(project["id"], user_id)
        if evidence_table is None:
            evidence_table = await self._build_context_paper_table(context)
        topic = context.get("research_project_name") or project.get("title") or "当前写作项目"
        venue = " ".join(str(item) for item in [context.get("target_venue"), context.get("target_year")] if item).strip() or "待定投稿目标"
        collections = "、".join(context.get("collection_names") or []) or "未绑定论文分类"
        paper_count = len(context.get("paper_ids") or [])
        description = context.get("description") or project.get("description") or "暂无研究方向描述。"
        table_md = evidence_table.get("markdown") if isinstance(evidence_table, dict) else str(evidence_table or "")
        references = await self._build_context_reference_seed(context)
        return {
            "Introduction": (
                f"写作主题：{topic}\n\n"
                f"投稿目标：{venue}\n\n"
                f"研究背景：{description}\n\n"
                "本章节是基于绑定研究方向和论文分类生成的结构化起点。正式写作时应进一步补充问题定义、贡献列表和实验结论。"
            ),
            "Related Work": (
                f"当前草稿绑定了论文分类：{collections}，初始证据论文 {paper_count} 篇。\n\n"
                "建议先根据下方证据对比表梳理支持证据、基线方法和反例，再逐段扩写 Related Work。"
            ),
            "Related Work Comparison Table": table_md,
            "References": references,
        }

    async def _build_context_paper_table(self, context: dict) -> dict:
        from app.db.models.paper import Paper

        paper_ids = context.get("paper_ids") or []
        if not paper_ids:
            return {"markdown": "暂无绑定论文。建议先选择论文分类或从论文库补充证据。"}
        rows = []
        for index, paper_id in enumerate(paper_ids[:12], 1):
            try:
                result = await self.session.execute(select(Paper).where(Paper.id == UUID(paper_id)))
                paper = result.scalar_one_or_none()
            except ValueError:
                paper = None
            if not paper:
                continue
            rows.append(
                "| {idx} | {title} | {year} | {source} | {hint} |".format(
                    idx=index,
                    title=self._escape_table_cell(paper.title),
                    year=paper.year or "N/A",
                    source=self._escape_table_cell(paper.source or "local"),
                    hint=self._escape_table_cell((paper.abstract or "待补全文后判断贡献")[:160]),
                )
            )
        header = "| # | 论文 | 年份 | 来源 | 初始写作提示 |\n|---|---|---:|---|---|"
        return {"markdown": "\n".join(["### Context-bound Evidence Table", "", header, *rows])}

    async def _build_context_reference_seed(self, context: dict) -> str:
        from app.db.models.paper import Paper

        refs = []
        for index, paper_id in enumerate(context.get("paper_ids") or [], 1):
            try:
                result = await self.session.execute(select(Paper).where(Paper.id == UUID(paper_id)))
                paper = result.scalar_one_or_none()
            except ValueError:
                paper = None
            if not paper:
                continue
            identifiers = [f"Paper ID: {paper.id}"]
            if paper.arxiv_id:
                identifiers.append(f"arXiv: {paper.arxiv_id}")
            if paper.doi:
                identifiers.append(f"DOI: {paper.doi}")
            refs.append(f"[{index}] {paper.title} ({'; '.join(identifiers)})")
        return "\n".join(refs) or "暂无可引用论文。请先绑定论文分类或补充证据。"

    async def create_review_draft_from_topic(
        self,
        user_id: str,
        topic: str,
        max_papers: int = 8,
        language: str = "chinese",
    ) -> dict:
        """从研究方向创建综述草稿项目。"""
        from app.services.writing_service import WritingAssistantService

        writing = WritingAssistantService(self.session)
        table = await writing.generate_related_work_table(topic, max_papers=max_papers)
        rows = table.get("rows", [])
        paper_ids = [row["paper_id"] for row in rows if row.get("paper_id")]

        metadata = {
            "source": "topic_review_draft",
            "source_topic": topic,
            "language": language,
            "recommended_paper_ids": paper_ids,
            "recommended_arxiv_ids": [],
            "related_work_table": table.get("markdown", ""),
            "evidence_status": "sufficient" if rows else "insufficient",
        }
        project = await self.create_project(
            user_id=user_id,
            title=f"{topic} 综述草稿",
            description=f"由研究方向「{topic}」自动生成的综述初稿，基于本地论文库检索结果。",
            template_type="survey",
            metadata_json=metadata,
        )

        sections_by_title = {section["title"]: section for section in project.get("sections", [])}
        content_by_title = self._build_review_draft_sections(topic, rows, table)
        for title, content in content_by_title.items():
            section = sections_by_title.get(title)
            if section:
                await self.update_section(
                    section["id"],
                    user_id,
                    content=content,
                    status="draft",
                )

        updated = await self.get_project(project["id"], user_id)
        return {
            "project": updated,
            "related_work_table": table,
            "evidence_status": metadata["evidence_status"],
        }

    async def create_review_draft_from_research_idea(
        self,
        user_id: str,
        research_project,
        idea,
        writing_brief: Optional[dict] = None,
    ) -> dict:
        """从证据驱动 Research Idea 创建写作草稿。"""
        evidence_items = list((idea.evidence_json or {}).get("items") or [])
        local_paper_ids = self._extract_local_paper_ids(evidence_items)
        writing_brief = writing_brief or self.build_research_idea_writing_brief(research_project, idea)
        metadata = {
            "source": "research_idea",
            "source_project_id": str(research_project.id),
            "source_project_name": research_project.name,
            "source_idea_id": str(idea.id),
            "source_idea_title": idea.title,
            "recommended_paper_ids": local_paper_ids,
            "recommended_arxiv_ids": [
                item.get("arxiv_id") for item in evidence_items if item.get("arxiv_id")
            ],
            "evidence_items": evidence_items,
            "evidence_status": "sufficient" if local_paper_ids or evidence_items else "insufficient",
            "writing_brief": writing_brief,
            "claim_evidence_map": writing_brief.get("claim_evidence_map") or [],
            "unsafe_claims": writing_brief.get("unsafe_claims") or [],
        }
        project = await self.create_project(
            user_id=user_id,
            title=f"{idea.title} 写作草稿",
            description=f"由研究方向「{research_project.name}」中的 Proposal 自动创建。",
            template_type="survey",
            metadata_json=metadata,
        )

        sections_by_title = {section["title"]: section for section in project.get("sections", [])}
        content_by_title = self._build_research_idea_draft_sections(
            research_project,
            idea,
            evidence_items,
            writing_brief=writing_brief,
        )
        for title, content in content_by_title.items():
            section = sections_by_title.get(title)
            if section:
                await self.update_section(
                    section["id"],
                    user_id,
                    content=content,
                    status="draft",
                )

        updated = await self.get_project(project["id"], user_id)
        return {
            "project": updated,
            "evidence_status": metadata["evidence_status"],
            "evidence_count": len(evidence_items),
            "local_paper_count": len(local_paper_ids),
            "writing_brief": writing_brief,
        }

    def build_research_idea_writing_brief(self, research_project, idea) -> dict:
        """Build a bounded, citation-conservative writing brief for a Proposal."""
        from app.services.research_idea_workbench import ResearchIdeaWorkbenchService

        evidence_items = [item for item in list((idea.evidence_json or {}).get("items") or []) if isinstance(item, dict)]
        local_paper_ids = self._extract_local_paper_ids(evidence_items)
        workbench = ResearchIdeaWorkbenchService(self.session)
        if not hasattr(idea, "referenced_papers"):
            idea.referenced_papers = {
                "paper_ids": [
                    item.get("paper_id")
                    for item in evidence_items
                    if item.get("paper_id")
                ]
            }
        if not hasattr(idea, "feasibility_score"):
            idea.feasibility_score = None
        if not hasattr(idea, "novelty_score"):
            idea.novelty_score = None
        if not hasattr(idea, "project_id"):
            idea.project_id = getattr(research_project, "id", "")
        validation = workbench.validate_idea(idea, research_project)
        execution = workbench.build_experiment_execution_pack(idea, research_project)
        review = idea.review_json or {}
        proposal_review = review.get("proposal_review") if isinstance(review.get("proposal_review"), dict) else {}
        plan = idea.experiment_plan or {}
        evidence_refs = [self._writing_brief_evidence_ref(item, index) for index, item in enumerate(evidence_items, start=1)]

        claims = self._writing_brief_claims(idea, proposal_review)
        claim_map = self._build_claim_evidence_map(claims, evidence_refs)
        unsupported_claims = [item["claim"] for item in claim_map if item["status"] == "unsupported"]
        evidence_gaps = self._writing_brief_evidence_gaps(validation, execution, proposal_review, evidence_items)
        unsafe_claims = list(dict.fromkeys([
            *unsupported_claims,
            *self._brief_list((validation.get("writing_readiness") or {}).get("reasons"), limit=4),
            *self._brief_list(proposal_review.get("reviewer_objections"), limit=4),
        ]))[:8]

        title_candidates = list(dict.fromkeys([
            str(idea.title or "Untitled Proposal").strip(),
            f"{idea.title}: Evidence-Grounded Study" if idea.title else "",
            f"{getattr(research_project, 'name', 'Research')} via {idea.title}" if idea.title else "",
        ]))
        title_candidates = [item[:180] for item in title_candidates if item][:3]

        section_outline = self._build_writing_brief_outline(idea, evidence_refs, proposal_review, plan)
        experiment_plan = self._brief_list(plan.get("steps"), limit=6)
        if not experiment_plan:
            experiment_plan = self._brief_list(execution.get("next_actions"), limit=6)

        return {
            "idea_id": str(idea.id),
            "project_id": str(getattr(idea, "project_id", getattr(research_project, "id", ""))),
            "source_project_name": getattr(research_project, "name", ""),
            "title_candidates": title_candidates,
            "abstract_draft": self._writing_brief_abstract(research_project, idea, proposal_review, plan),
            "contribution_chain": self._build_contribution_chain(idea, proposal_review, claim_map),
            "section_outline": section_outline,
            "claim_evidence_map": claim_map,
            "evidence_gaps": evidence_gaps,
            "experiment_writing_plan": experiment_plan[:6],
            "limitations": self._writing_brief_limitations(validation, proposal_review, execution),
            "unsafe_claims": unsafe_claims,
            "evidence_status": "sufficient" if local_paper_ids or evidence_items else "insufficient",
            "evidence_count": len(evidence_items),
            "local_paper_count": len(local_paper_ids),
            "review_readiness": proposal_review.get("writing_readiness") or (validation.get("writing_readiness") or {}).get("status"),
        }

    def _extract_local_paper_ids(self, evidence_items: list[dict]) -> list[str]:
        ids: list[str] = []
        for item in evidence_items:
            candidate = item.get("imported_paper_id") or item.get("paper_id")
            if candidate and self._looks_like_uuid(str(candidate)) and str(candidate) not in ids:
                ids.append(str(candidate))
        return ids

    def _looks_like_uuid(self, value: str) -> bool:
        return bool(re.fullmatch(r"[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12}", value or ""))

    def _build_research_idea_draft_sections(
        self,
        research_project,
        idea,
        evidence_items: list[dict],
        *,
        writing_brief: Optional[dict] = None,
    ) -> dict:
        writing_brief = writing_brief or self.build_research_idea_writing_brief(research_project, idea)
        evidence_table = self._build_evidence_markdown_table(evidence_items)
        references = self._build_evidence_references(evidence_items)
        if not evidence_items:
            evidence_note = "当前 Proposal 暂无可用证据论文。请先运行 Idea 工作台或导入外部证据，再继续完善 Related Work。"
        else:
            evidence_note = f"当前草稿继承了 {len(evidence_items)} 条 Proposal 证据，其中本地论文可用于后续 BibTeX 导出。"

        review = idea.review_json or {}
        plan = idea.experiment_plan or {}
        scores = review.get("scores") or {}
        score_text = "、".join(f"{key}: {value}/10" for key, value in scores.items()) or "暂无六维评分"
        baselines = "、".join(plan.get("baselines") or []) or "待补充"
        metrics = "、".join(plan.get("metrics") or []) or "待补充"
        steps = "\n".join(f"{index + 1}. {step}" for index, step in enumerate(plan.get("steps") or [])) or "待补充最小实验步骤。"
        claim_lines = "\n".join(
            f"- [{item.get('status')}] {item.get('claim')} -> {', '.join(ref.get('marker', '') for ref in item.get('evidence_refs', []) if ref.get('marker')) or '暂无证据'}"
            for item in writing_brief.get("claim_evidence_map", [])[:6]
        ) or "- 暂无 claim-evidence 映射。"
        unsafe_lines = "\n".join(f"- {item}" for item in writing_brief.get("unsafe_claims", [])[:6]) or "- 暂无明确不可写 claim。"
        outline_lines = "\n".join(
            f"- **{item.get('section')}**：{item.get('purpose')} {item.get('seed_content') or ''}"
            for item in writing_brief.get("section_outline", [])[:8]
        ) or "- 暂无章节骨架。"

        related_work = (
            f"{evidence_note}\n\n"
            "下面的证据表按 Proposal 生成时的证据角色保留，写作时建议逐条确认其是否能够支撑对应论断。\n\n"
            f"{evidence_table}\n\n"
            "## Claim-Evidence Map\n\n"
            f"{claim_lines}"
        )
        return {
            "Abstract": writing_brief.get("abstract_draft") or (
                f"本文围绕「{research_project.name}」中的 Proposal「{idea.title}」展开写作。"
                f"核心假设是：{idea.hypothesis or '待补充可证伪假设'}。"
                "当前草稿将 Proposal 的证据、方法定位和最小实验计划转换为可编辑写作结构。"
            ),
            "Introduction": (
                f"研究方向：{research_project.name}\n\n"
                f"{research_project.description or '暂无研究方向描述。'}\n\n"
                f"Proposal 描述：{idea.description or '暂无描述。'}\n\n"
                f"技术草图：{idea.approach or '待补充技术路线。'}\n\n"
                "## 写作章节骨架\n\n"
                f"{outline_lines}"
            ),
            "Related Work": related_work,
            "Related Work Comparison Table": evidence_table,
            "Research Gaps": (
                f"创新点：{idea.novelty or '待补充'}\n\n"
                f"评审理由：{review.get('rationale') or '暂无评审理由'}\n\n"
                f"主要不确定性：{review.get('uncertainty') or '暂无'}\n\n"
                f"六维评分：{score_text}\n\n"
                f"最小实验数据集：{plan.get('dataset') or '待补充'}\n\n"
                f"基线方法：{baselines}\n\n"
                f"评测指标：{metrics}\n\n"
                f"实验步骤：\n{steps}\n\n"
                "## 暂不应直接写成结论的 Claim\n\n"
                f"{unsafe_lines}"
            ),
            "References": references or "暂无引用。建议先导入或补全文证据论文。",
        }

    def _writing_brief_evidence_ref(self, item: dict, index: int) -> dict:
        local_id = item.get("imported_paper_id") or item.get("paper_id")
        return {
            "marker": f"[{index}]",
            "paper_id": str(item.get("paper_id") or ""),
            "local_paper_id": str(local_id) if local_id and self._looks_like_uuid(str(local_id)) else None,
            "title": item.get("title") or "Untitled evidence",
            "year": item.get("year"),
            "role": item.get("category") or "evidence",
            "relevance": item.get("relevance") or item.get("abstract_excerpt") or item.get("source") or "",
            "score": item.get("score"),
            "arxiv_id": item.get("arxiv_id"),
            "doi": item.get("doi"),
        }

    def _writing_brief_claims(self, idea, proposal_review: dict) -> list[str]:
        return list(dict.fromkeys([
            str(idea.hypothesis or "").strip(),
            str(idea.novelty or "").strip(),
            str(idea.approach or "").strip(),
            *self._brief_list(proposal_review.get("contributions"), limit=4),
        ]))[:8]

    def _build_claim_evidence_map(self, claims: list[str], evidence_refs: list[dict]) -> list[dict]:
        claim_map = []
        for index, claim in enumerate([item for item in claims if item], start=1):
            refs = evidence_refs[:3] if evidence_refs else []
            status = "supported" if len(refs) >= 2 else "partially_supported" if refs else "unsupported"
            claim_map.append({
                "claim": claim[:360],
                "status": status,
                "evidence_refs": refs,
                "writing_use": "可作为正文论断起点，定稿前仍需核对全文证据。" if refs else "证据不足，暂时只能写为待验证假设或未来工作。",
                "priority": index,
            })
        if not claim_map:
            claim_map.append({
                "claim": "Proposal 尚未形成明确 claim。",
                "status": "unsupported",
                "evidence_refs": [],
                "writing_use": "先补充可证伪假设和证据。",
                "priority": 1,
            })
        return claim_map

    def _build_contribution_chain(self, idea, proposal_review: dict, claim_map: list[dict]) -> list[dict]:
        contributions = self._brief_list(proposal_review.get("contributions"), limit=4)
        if not contributions:
            contributions = [str(idea.novelty or idea.hypothesis or "待明确贡献点")[:240]]
        return [
            {
                "step": index,
                "claim": contribution,
                "evidence_status": claim_map[min(index - 1, len(claim_map) - 1)].get("status") if claim_map else "unsupported",
                "writing_goal": "把该贡献写成可验证、可被证据支撑的论文贡献句。",
            }
            for index, contribution in enumerate(contributions[:4], start=1)
        ]

    def _build_writing_brief_outline(self, idea, evidence_refs: list[dict], proposal_review: dict, plan: dict) -> list[dict]:
        return [
            {"section": "Abstract", "purpose": "压缩问题、假设、方法和证据状态。", "seed_content": str(idea.hypothesis or "")[:240], "evidence_ids": []},
            {"section": "Introduction", "purpose": "建立研究问题和贡献链。", "seed_content": str(idea.description or "")[:240], "evidence_ids": [ref["paper_id"] for ref in evidence_refs[:3]]},
            {"section": "Related Work", "purpose": "用现有证据组织相似工作和差异点。", "seed_content": proposal_review.get("summary") or "", "evidence_ids": [ref["paper_id"] for ref in evidence_refs]},
            {"section": "Method", "purpose": "解释技术路线和关键假设。", "seed_content": str(idea.approach or "")[:240], "evidence_ids": []},
            {"section": "Experiments", "purpose": "写清数据集、强基线、指标、消融和失败判定。", "seed_content": str(plan.get("dataset") or "")[:180], "evidence_ids": []},
            {"section": "Limitations", "purpose": "提前保留审稿风险和不可过度宣称的点。", "seed_content": "；".join(self._brief_list(proposal_review.get("reviewer_objections"), limit=3)), "evidence_ids": []},
        ]

    def _writing_brief_abstract(self, research_project, idea, proposal_review: dict, plan: dict) -> str:
        readiness = proposal_review.get("writing_readiness")
        readiness_note = f"当前审稿式写作准备状态为 {readiness}。" if readiness else ""
        dataset = plan.get("dataset") or "待补充数据集"
        return (
            f"本文围绕「{getattr(research_project, 'name', '研究方向')}」中的 Proposal「{idea.title}」展开。"
            f"核心假设是：{idea.hypothesis or '待补充可证伪假设'}。"
            f"方法上，草稿将以「{idea.approach or '待补充技术路线'}」为主线，并围绕 {dataset} 设计最小实验。"
            f"{readiness_note} 定稿前需要逐条核对 claim 与证据引用。"
        )

    def _writing_brief_evidence_gaps(self, validation: dict, execution: dict, proposal_review: dict, evidence_items: list[dict]) -> list[str]:
        gaps = []
        if not evidence_items:
            gaps.append("缺少 Proposal 证据，Related Work 和贡献论证暂不能定稿。")
        coverage = validation.get("coverage") or {}
        if not coverage.get("has_enough_evidence"):
            gaps.append("证据覆盖不足，至少补充 2 篇可核对论文后再写强结论。")
        gaps.extend(self._brief_list(execution.get("next_actions"), limit=3))
        gaps.extend(self._brief_list(proposal_review.get("required_experiments"), limit=3))
        return list(dict.fromkeys(gaps))[:8]

    def _writing_brief_limitations(self, validation: dict, proposal_review: dict, execution: dict) -> list[str]:
        limitations = self._brief_list(proposal_review.get("weakest_assumptions"), limit=4)
        limitations.extend(self._brief_list(proposal_review.get("reviewer_objections"), limit=4))
        limitations.extend(self._brief_list(validation.get("feasibility_risks"), limit=4))
        limitations.extend(self._brief_list(execution.get("risks"), limit=4))
        return list(dict.fromkeys(limitations))[:8]

    def _brief_list(self, value: Any, *, limit: int = 6) -> list[str]:
        if value is None:
            return []
        if isinstance(value, list):
            items = value
        else:
            items = [value]
        result = []
        for item in items:
            if isinstance(item, dict):
                text = item.get("message") or item.get("detail") or item.get("label") or item.get("title") or item.get("claim")
            else:
                text = item
            text = str(text or "").strip()
            if text and text not in result:
                result.append(text[:360])
            if len(result) >= limit:
                break
        return result

    def _build_evidence_markdown_table(self, evidence_items: list[dict]) -> str:
        header = "| # | 证据论文 | 年份 | 角色 | 相关性 | 可用信息 |\n|---|---|---:|---|---:|---|"
        if not evidence_items:
            return f"{header}\n| - | 暂无证据 | - | - | - | 建议先运行 Idea 工作台或导入论文 |"
        rows = []
        for index, item in enumerate(evidence_items, start=1):
            title = self._escape_table_cell(item.get("title") or "Untitled")
            year = item.get("year") or "N/A"
            category = item.get("category") or "evidence"
            score = item.get("score")
            score_text = f"{float(score):.2f}" if isinstance(score, (int, float)) else "N/A"
            info = item.get("relevance") or item.get("abstract_excerpt") or item.get("source") or "待补充摘要/全文"
            rows.append(f"| {index} | {title} | {year} | {category} | {score_text} | {self._escape_table_cell(str(info)[:180])} |")
        return f"{header}\n" + "\n".join(rows)

    def _build_evidence_references(self, evidence_items: list[dict]) -> str:
        refs = []
        for index, item in enumerate(evidence_items, start=1):
            identifiers = []
            local_id = item.get("imported_paper_id") or item.get("paper_id")
            if local_id and self._looks_like_uuid(str(local_id)):
                identifiers.append(f"Paper ID: {local_id}")
            if item.get("arxiv_id"):
                identifiers.append(f"arXiv: {item['arxiv_id']}")
            if item.get("doi"):
                identifiers.append(f"DOI: {item['doi']}")
            suffix = f" ({'; '.join(identifiers)})" if identifiers else ""
            refs.append(f"[{index}] {item.get('title') or 'Untitled evidence'}{suffix}")
        return "\n".join(refs)

    def _escape_table_cell(self, value: str) -> str:
        return (value or "").replace("\n", " ").replace("|", "\\|")

    def _build_review_draft_sections(self, topic: str, rows: list[dict], table: dict) -> dict:
        if not rows:
            insufficient = (
                f"当前本地知识库暂未检索到与「{topic}」足够相关的论文。\n\n"
                "建议先在论文库中检索外部论文并入库，随后补全文与向量，再重新生成综述草稿。"
            )
            return {
                "Abstract": insufficient,
                "Introduction": insufficient,
                "Related Work": insufficient,
                "Related Work Comparison Table": table.get("markdown", ""),
                "Research Gaps": "证据不足，暂不生成研究空白判断。",
                "References": "暂无可引用论文。",
            }

        citations = ", ".join(f"[{row['index']}]" for row in rows[:5])
        top_titles = "、".join(row["title"] for row in rows[:3])
        related_lines = [
            f"- [{row['index']}] **{row['title']}** ({row.get('year') or 'N/A'}): {row.get('contribution', '')} 该工作可作为「{row.get('role_label')}」使用，{row.get('comparison_point', '')}"
            for row in rows
        ]
        gap_lines = [
            "- 若多篇论文集中在相同数据集或设定上，后续综述需要补充跨数据集、跨场景或真实应用中的证据。",
            "- 对摘要信息不足的论文，建议补全文后再判断其方法细节、实验指标和局限性。",
            "- 写作定稿前应逐句检查引用是否能支撑对应结论，避免只凭题名或摘要推测。",
        ]
        refs = [
            f"[{row['index']}] {row['title']} ({row.get('year') or 'N/A'}). Paper ID: {row['paper_id']}"
            for row in rows
        ]
        return {
            "Abstract": (
                f"本文围绕「{topic}」进行文献综述，初步整理了本地知识库中 {len(rows)} 篇相关论文。"
                f"草稿重点关注代表性工作（如 {top_titles}）之间的方法差异、证据角色和可对比点。"
                "当前版本适合作为写作起点，定稿前仍需结合全文证据补充更细粒度的实验和局限分析。"
            ),
            "Introduction": (
                f"「{topic}」是当前研究中值得系统梳理的方向。已有工作从不同方法路线、实验设置和应用场景展开探索，"
                f"其中 {citations} 可作为初步参考。为了避免简单罗列论文，本文将相关工作按贡献、证据角色和可对比点组织，"
                "并在后续章节进一步讨论仍需补足的研究空白。"
            ),
            "Related Work": "\n".join(related_lines),
            "Related Work Comparison Table": table.get("markdown", ""),
            "Research Gaps": "\n".join(gap_lines),
            "References": "\n".join(refs),
        }

    async def list_projects(self, user_id: str) -> list:
        """列出用户的所有项目。"""
        from app.db.models.writing import WritingProject, WritingSection
        from sqlalchemy.orm import selectinload

        result = await self.session.execute(
            select(WritingProject)
            .where(WritingProject.user_id == user_id)
            .options(selectinload(WritingProject.sections))
            .order_by(WritingProject.updated_at.desc())
        )
        projects = result.scalars().all()
        return [self._project_to_dict(p) for p in projects]

    async def get_project(self, project_id: str, user_id: str) -> dict | None:
        """获取项目详情。"""
        from app.db.models.writing import WritingProject
        from sqlalchemy.orm import selectinload
        from app.services.workspace_service import WorkspaceService

        try:
            pid = UUID(project_id)
        except ValueError:
            return None

        result = await self.session.execute(
            select(WritingProject)
            .where(WritingProject.id == pid)
            .options(selectinload(WritingProject.sections))
        )
        project = result.scalar_one_or_none()
        if not project:
            return None
        if str(project.user_id) != str(user_id):
            role = await WorkspaceService(self.session).resource_role_for_user(user_id, "writing_projects", str(project.id))
            if not WorkspaceService(self.session).role_can_read_resource(role):
                return None
            data = self._project_to_dict(project)
            data["workspace_access_role"] = role
            return data
        return self._project_to_dict(project)

    async def update_project(self, project_id: str, user_id: str,
                              **kwargs) -> dict | None:
        """更新项目。"""
        from app.db.models.writing import WritingProject
        from app.services.workspace_service import WorkspaceService

        try:
            pid = UUID(project_id)
        except ValueError:
            return None

        result = await self.session.execute(
            select(WritingProject)
            .where(WritingProject.id == pid)
        )
        project = result.scalar_one_or_none()
        if not project:
            return None
        if str(project.user_id) != str(user_id):
            workspace = WorkspaceService(self.session)
            role = await workspace.resource_role_for_user(user_id, "writing_projects", str(project.id))
            if not workspace.role_can_edit_resource(role):
                return None

        for key, value in kwargs.items():
            if hasattr(project, key):
                setattr(project, key, value)

        await self.session.commit()
        await self.session.refresh(project)
        return self._project_to_dict(project)

    async def bind_submission_profile(
        self,
        project_id: str,
        user_id: str,
        venue: str,
        year: str,
        template_inspection: dict,
    ) -> dict | None:
        """Bind target venue/year and inspected template metadata to a writing project."""
        project = await self.get_project(project_id, user_id)
        if not project:
            return None
        metadata = dict(project.get("metadata_json") or {})
        inspection = dict(template_inspection or {})
        profile = {
            "venue": (venue or "").strip(),
            "year": (year or "").strip(),
            "template_source": inspection.get("source_filename") or "",
            "template_status": inspection.get("status") or "needs_review",
            "status_label": inspection.get("status_label") or "需要人工确认",
            "document_class": inspection.get("document_class") or "",
            "main_tex": inspection.get("main_tex") or "",
            "class_files": inspection.get("class_files") or [],
            "style_files": inspection.get("style_files") or [],
            "packages": inspection.get("packages") or [],
            "venue_hints": inspection.get("venue_hints") or [],
            "warnings": inspection.get("warnings") or [],
        }
        if not profile["venue"] and profile["venue_hints"]:
            profile["venue"] = profile["venue_hints"][0]
        if not profile["venue"]:
            profile["warnings"] = [*profile["warnings"], "尚未填写目标会议/期刊。"]
        metadata["submission_profile"] = profile
        return await self.update_project(project_id, user_id, metadata_json=metadata)

    async def delete_project(self, project_id: str, user_id: str) -> bool:
        """删除项目。"""
        from app.db.models.writing import WritingProject

        try:
            pid = UUID(project_id)
        except ValueError:
            return False

        result = await self.session.execute(
            select(WritingProject)
            .where(WritingProject.id == pid, WritingProject.user_id == user_id)
        )
        project = result.scalar_one_or_none()
        if not project:
            return False

        await self.session.delete(project)
        await self.session.commit()
        return True

    # --- 章节管理 ---

    async def update_section(self, section_id: str, user_id: str,
                              **kwargs) -> dict | None:
        """更新章节内容/状态/标题。"""
        from app.db.models.writing import WritingSection, WritingProject
        from app.services.workspace_service import WorkspaceService

        try:
            sid = UUID(section_id)
        except ValueError:
            return None

        result = await self.session.execute(
            select(WritingSection)
            .join(WritingProject, WritingSection.project_id == WritingProject.id)
            .where(WritingSection.id == sid)
        )
        section = result.scalar_one_or_none()
        if not section:
            return None
        project_result = await self.session.execute(select(WritingProject).where(WritingProject.id == section.project_id))
        project = project_result.scalar_one_or_none()
        if not project:
            return None
        if str(project.user_id) != str(user_id):
            workspace = WorkspaceService(self.session)
            role = await workspace.resource_role_for_user(user_id, "writing_projects", str(project.id))
            if not workspace.role_can_edit_resource(role):
                return None

        for key, value in kwargs.items():
            if hasattr(section, key):
                setattr(section, key, value)

        # 自动计算字数
        if "content" in kwargs:
            section.word_count = len(kwargs["content"])

        await self.session.commit()
        await self.session.refresh(section)
        return self._section_to_dict(section)

    async def reorder_sections(self, project_id: str, user_id: str,
                                section_ids: List[str]) -> bool:
        """重排章节顺序。"""
        from app.db.models.writing import WritingProject, WritingSection
        from app.services.workspace_service import WorkspaceService

        try:
            pid = UUID(project_id)
        except ValueError:
            return False

        result = await self.session.execute(select(WritingProject).where(WritingProject.id == pid))
        project = result.scalar_one_or_none()
        if not project:
            return False
        if str(project.user_id) != str(user_id):
            workspace = WorkspaceService(self.session)
            role = await workspace.resource_role_for_user(user_id, "writing_projects", str(project.id))
            if not workspace.role_can_edit_resource(role):
                return False

        for i, sid in enumerate(section_ids):
            try:
                await self.session.execute(
                    update(WritingSection)
                    .where(WritingSection.id == UUID(sid), WritingSection.project_id == str(pid))
                    .values(order=i)
                )
            except (ValueError, Exception):
                continue

        await self.session.commit()
        return True

    # --- 证据卡片与引用校验 ---

    async def get_evidence_cards(self, project_id: str, user_id: str) -> dict | None:
        """获取写作项目关联证据卡片。"""
        from app.db.models.paper import Paper

        project = await self.get_project(project_id, user_id)
        if not project:
            return None

        metadata = project.get("metadata_json") or {}
        evidence_items = list(metadata.get("evidence_items") or [])
        recommended_ids = list(dict.fromkeys(metadata.get("recommended_paper_ids") or []))
        recommended_arxiv_ids = list(dict.fromkeys(metadata.get("recommended_arxiv_ids") or []))

        local_ids = list(recommended_ids)
        for item in evidence_items:
            candidate = item.get("imported_paper_id") or item.get("paper_id")
            if candidate and self._looks_like_uuid(str(candidate)) and str(candidate) not in local_ids:
                local_ids.append(str(candidate))

        local_papers: dict[str, Paper] = {}
        for paper_id in local_ids:
            try:
                result = await self.session.execute(select(Paper).where(Paper.id == UUID(paper_id)))
                paper = result.scalar_one_or_none()
                if paper:
                    local_papers[str(paper.id)] = paper
            except ValueError:
                continue

        cards: list[dict[str, Any]] = []
        seen: set[str] = set()
        for item in evidence_items:
            local_id = item.get("imported_paper_id") or item.get("paper_id")
            paper = local_papers.get(str(local_id)) if local_id else None
            card = self._evidence_item_to_card(item, paper, len(cards) + 1)
            key = self._evidence_card_key(card)
            if key not in seen:
                seen.add(key)
                cards.append(card)

        if not evidence_items:
            for paper_id in recommended_ids:
                paper = local_papers.get(str(paper_id))
                if not paper:
                    continue
                card = self._paper_to_evidence_card(paper, len(cards) + 1)
                key = self._evidence_card_key(card)
                if key not in seen:
                    seen.add(key)
                    cards.append(card)

        for arxiv_id in recommended_arxiv_ids:
            if not arxiv_id:
                continue
            normalized = self._normalize_arxiv_id(str(arxiv_id))
            if any(self._normalize_arxiv_id(card.get("arxiv_id") or "") == normalized for card in cards):
                continue
            card = {
                "id": f"arxiv:{normalized}",
                "index": len(cards) + 1,
                "citation_marker": f"[{len(cards) + 1}]",
                "title": f"arXiv:{arxiv_id}",
                "year": None,
                "authors": "",
                "source": "external",
                "source_label": "外部 arXiv",
                "source_kind": "external",
                "paper_id": None,
                "external_paper_id": None,
                "arxiv_id": arxiv_id,
                "doi": None,
                "role": "supporting_evidence",
                "role_label": "支持证据",
                "snippet": "该论文尚未导入本地论文库，无法进行引用支撑强度校验。",
                "local_status": "external",
                "local_status_label": "未入库",
                "bibtex_ready": False,
            }
            seen.add(self._evidence_card_key(card))
            cards.append(card)

        coverage = {
            "total": len(cards),
            "local": sum(1 for card in cards if card["local_status"] == "local"),
            "external": sum(1 for card in cards if card["local_status"] == "external"),
            "bibtex_ready": sum(1 for card in cards if card.get("bibtex_ready")),
        }
        return {
            "project_id": project["id"],
            "source": metadata.get("source") or "manual",
            "evidence_status": metadata.get("evidence_status") or ("sufficient" if cards else "insufficient"),
            "coverage": coverage,
            "cards": cards,
        }

    async def check_section_citations(
        self,
        project_id: str,
        user_id: str,
        text: str,
        section_id: Optional[str] = None,
    ) -> dict | None:
        """检查章节中的引用是否能被项目证据支撑。"""
        from app.db.models.paper import Paper
        from app.services.writing_service import WritingAssistantService

        evidence = await self.get_evidence_cards(project_id, user_id)
        if not evidence:
            return None

        cards = evidence.get("cards") or []
        citations = self._extract_citation_mentions(text or "")
        checks: list[dict[str, Any]] = []
        writer = WritingAssistantService(self.session)

        if not citations:
            return {
                "project_id": project_id,
                "section_id": section_id,
                "checks": [{
                    "citation": None,
                    "status": "missing",
                    "label": "缺少引用",
                    "sentence": (text or "").strip()[:240],
                    "explanation": "该章节没有检测到 [1]、arXiv 或 Paper ID 引用标记。建议为关键结论补充证据引用。",
                    "decision_action": "先为关键结论插入证据卡引用",
                    "decision_warning": "没有引用标记时，导出前引用校验无法判断正文结论是否有证据支持。",
                    "card": None,
                }],
                "summary": self._citation_summary([{"status": "missing"}]),
                "recommendations": self._recommend_evidence_cards(cards),
            }

        for mention in citations:
            card = self._resolve_citation_card(mention, cards)
            sentence = self._citation_context(text or "", mention["start"], mention["end"])
            if not card:
                checks.append({
                    "citation": mention["raw"],
                    "status": "missing",
                    "label": "未找到证据卡",
                    "sentence": sentence,
                    "explanation": "该引用没有匹配到当前写作项目的证据卡，请检查编号或补充 References。",
                    "decision_action": "检查引用编号或补充证据卡",
                    "decision_warning": "未匹配到证据卡的引用不应直接进入终稿。",
                    "card": None,
                })
                continue

            if card.get("local_status") == "unknown":
                checks.append({
                    "citation": mention["raw"],
                    "status": "missing",
                    "label": "本地记录缺失",
                    "sentence": sentence,
                    "explanation": "证据卡保存了本地论文 ID，但当前数据库没有找到该论文。建议重新入库或替换引用。",
                    "decision_action": "重新入库该论文或替换引用",
                    "decision_warning": "本地记录缺失时无法校验引用是否真实支撑当前句子。",
                    "card": card,
                })
                continue

            if card.get("local_status") != "local" or not card.get("paper_id"):
                checks.append({
                    "citation": mention["raw"],
                    "status": "unchecked",
                    "label": "外部证据未校验",
                    "sentence": sentence,
                    "explanation": "该证据尚未导入本地论文库，系统无法判断它是否支撑当前句子。建议先一键入库并补全文/补向量。",
                    "decision_action": "先入库并补全文/补向量",
                    "decision_warning": "外部证据未校验时，只能作为临时占位引用。",
                    "card": card,
                })
                continue

            try:
                result = await self.session.execute(select(Paper).where(Paper.id == UUID(card["paper_id"])))
                paper = result.scalar_one_or_none()
            except ValueError:
                paper = None

            if not paper:
                checks.append({
                    "citation": mention["raw"],
                    "status": "missing",
                    "label": "本地论文缺失",
                    "sentence": sentence,
                    "explanation": "证据卡记录了本地论文 ID，但当前数据库没有找到该论文。",
                    "decision_action": "修复证据卡或替换引用",
                    "decision_warning": "数据库中找不到论文时，引用真实性无法确认。",
                    "card": card,
                })
                continue

            match = writer.score_sentence_paper_match(sentence, paper)
            role = writer.classify_citation_role(sentence, paper)
            decision = writer.build_citation_decision(role, match)
            checks.append({
                "citation": mention["raw"],
                "status": match["match_status"],
                "label": match["match_label"],
                "sentence": sentence,
                "explanation": match["match_explanation"],
                "match_score": match["match_score"],
                "match_terms": match["match_terms"],
                "role": role["role"],
                "role_label": role["role_label"],
                **decision,
                "card": card,
            })

        return {
            "project_id": project_id,
            "section_id": section_id,
            "checks": checks,
            "summary": self._citation_summary(checks),
            "recommendations": self._recommend_evidence_cards(cards),
        }

    async def analyze_section_quality(
        self,
        project_id: str,
        user_id: str,
        title: str,
        text: str,
        section_id: Optional[str] = None,
    ) -> dict | None:
        """Deterministically evaluate whether a section is ready to polish/export."""
        project = await self.get_project(project_id, user_id)
        if not project:
            return None

        content = (text or "").strip()
        title = (title or "Untitled").strip()
        word_count = len(re.findall(r"[A-Za-z0-9_\-]+|[\u4e00-\u9fff]", content))
        paragraph_count = len([p for p in re.split(r"\n\s*\n", content) if p.strip()])
        sentence_count = len([s for s in re.split(r"[。！？.!?]\s*", content) if s.strip()])
        citations = self._extract_citation_mentions(content)
        lowered = content.lower()

        has_comparison = bool(re.search(r"baseline|benchmark|sota|compare|comparison|相比|对比|基线|已有工作|现有方法", lowered))
        has_gap = bool(re.search(r"gap|limitation|challenge|future|不足|局限|挑战|未解决|仍然|缺少|空白", lowered))
        has_evidence_language = bool(re.search(r"\[\d+\]|arxiv|paper id|实验|结果|数据集|evidence|benchmark|citation", lowered))

        dimensions = [
            self._quality_dimension(
                "claim",
                "核心论断",
                bool(content) and (word_count >= 40 or sentence_count >= 2),
                22 if word_count >= 80 else 16 if word_count >= 40 else 6 if content else 0,
                "章节已有可识别的核心表述。",
                "先用 1-2 句明确本节要证明或解释的核心 claim。",
            ),
            self._quality_dimension(
                "evidence",
                "证据支撑",
                bool(citations) or has_evidence_language,
                24 if len(citations) >= 2 else 18 if citations else 8 if has_evidence_language else 0,
                "章节包含引用或证据线索。",
                "为关键结论插入证据卡引用，或补充实验/数据集/原文片段。",
            ),
            self._quality_dimension(
                "comparison",
                "对比对象",
                has_comparison,
                18 if has_comparison else 5,
                "章节说明了与已有方法或基线的关系。",
                "补一句与 baseline、已有工作或相邻方法的差异。",
            ),
            self._quality_dimension(
                "gap",
                "研究空白",
                has_gap,
                18 if has_gap else 5,
                "章节点出了局限、挑战或研究空白。",
                "补一句说明现有工作还没有解决什么，以及本节如何过渡。",
            ),
            self._quality_dimension(
                "structure",
                "段落结构",
                paragraph_count >= 2 or sentence_count >= 4,
                18 if paragraph_count >= 2 and sentence_count >= 4 else 12 if sentence_count >= 3 else 4,
                "章节结构具备基本展开。",
                "把内容拆成“背景/证据/对比/小结”几个自然段。",
            ),
        ]

        score = min(100, sum(item["score"] for item in dimensions))
        if score >= 78:
            status = "ready"
            status_label = "可进入润色"
            summary = "章节已经具备较完整的论断、证据和结构，可以继续做语言润色与引用细查。"
        elif score >= 45:
            status = "needs_revision"
            status_label = "需要补强"
            summary = "章节有初稿基础，但仍需补充证据、对比或研究空白后再进入定稿。"
        else:
            status = "incomplete"
            status_label = "初稿不足"
            summary = "章节还没有形成可靠学术段落，建议先补核心论断和证据。"

        rewrite_actions = [
            {"key": item["key"], "label": item["label"], "action": item["rewrite_hint"]}
            for item in dimensions
            if item["status"] != "pass"
        ][:4]

        return {
            "project_id": project_id,
            "section_id": section_id,
            "title": title,
            "status": status,
            "status_label": status_label,
            "overall_score": score,
            "summary": summary,
            "metrics": {
                "word_count": word_count,
                "paragraph_count": paragraph_count,
                "sentence_count": sentence_count,
                "citation_count": len(citations),
            },
            "dimensions": dimensions,
            "rewrite_actions": rewrite_actions,
        }

    async def build_evidence_related_work_table(self, project_id: str, user_id: str) -> dict | None:
        """基于写作项目证据卡生成 Related Work 对比表。"""
        evidence = await self.get_evidence_cards(project_id, user_id)
        if not evidence:
            return None

        cards = evidence.get("cards") or []
        coverage = evidence.get("coverage") or {}
        header = (
            "| 引用 | 论文 | 年份 | 证据角色 | 状态 | 标识符 | 写作使用建议 |\n"
            "|---|---|---:|---|---|---|---|"
        )
        rows = []
        for card in cards:
            identifiers = []
            if card.get("arxiv_id"):
                identifiers.append(f"arXiv:{card['arxiv_id']}")
            if card.get("doi"):
                identifiers.append(f"DOI:{card['doi']}")
            if card.get("paper_id"):
                identifiers.append(f"Paper ID:{card['paper_id']}")
            rows.append(
                "| {marker} | {title} | {year} | {role} | {status} | {ids} | {use} |".format(
                    marker=self._escape_table_cell(card.get("citation_marker") or ""),
                    title=self._escape_table_cell(card.get("title") or "Untitled"),
                    year=card.get("year") or "N/A",
                    role=self._escape_table_cell(card.get("role_label") or "支持证据"),
                    status=self._escape_table_cell(card.get("local_status_label") or "未知"),
                    ids=self._escape_table_cell("；".join(identifiers) or "待补充"),
                    use=self._escape_table_cell(self._evidence_writing_use(card)),
                )
            )

        if not rows:
            rows.append("| - | 暂无证据 | - | - | 证据不足 | - | 请先从论文库入库或从研究方向生成证据。 |")

        warnings = []
        total = coverage.get("total") or len(cards)
        local = coverage.get("local") or 0
        external = coverage.get("external") or 0
        if total == 0:
            warnings.append("当前项目没有证据卡，无法生成可信 Related Work 对比表。")
        if total and local / total < 0.5:
            warnings.append("本地入库证据覆盖率偏低，表格中外部证据尚未经过全文/向量校验。")
        if external:
            warnings.append(f"有 {external} 条外部证据尚未入库，建议入库后再用于最终定稿。")

        markdown = "\n".join([
            "### Evidence-backed Related Work Comparison",
            "",
            header,
            *rows,
        ])
        if warnings:
            markdown += "\n\n> 证据提示：" + " ".join(warnings)

        return {
            "project_id": project_id,
            "markdown": markdown,
            "coverage": {
                "total": total,
                "local": local,
                "external": external,
                "bibtex_ready": coverage.get("bibtex_ready") or 0,
            },
            "warnings": warnings,
            "status": "weak_evidence" if warnings else "ready",
        }

    def _quality_dimension(
        self,
        key: str,
        label: str,
        passed: bool,
        score: int,
        pass_explanation: str,
        rewrite_hint: str,
    ) -> dict:
        if passed:
            status = "pass"
            explanation = pass_explanation
        elif score >= 10:
            status = "partial"
            explanation = f"{label}已有一些线索，但还不够明确。"
        else:
            status = "weak"
            explanation = f"{label}不足。"
        return {
            "key": key,
            "label": label,
            "status": status,
            "score": score,
            "explanation": explanation,
            "rewrite_hint": rewrite_hint,
        }

    # --- 导出 ---

    async def export_to_markdown(self, project_id: str, user_id: str) -> str | None:
        """导出为 Markdown。"""
        project = await self.get_project(project_id, user_id)
        if not project:
            return None

        lines = [f"# {project['title']}", "", f"*{project.get('description', '')}*", ""]
        for sec in project.get("sections", []):
            lines.append(f"## {sec['title']}")
            lines.append("")
            lines.append(sec.get("content", "") or "*（待撰写）*")
            lines.append("")

        return "\n".join(lines)

    async def export_to_latex(self, project_id: str, user_id: str) -> str | None:
        """导出为 LaTeX。"""
        project = await self.get_project(project_id, user_id)
        if not project:
            return None

        from app.services.latex_processor import latex_processor
        return latex_processor.render_to_tex(
            project["title"],
            project.get("sections", []),
            template="article",
        )

    async def export_to_docx_bytes(self, project_id: str, user_id: str):
        """导出为 Word (.docx) bytes。"""
        project = await self.get_project(project_id, user_id)
        if not project:
            return None

        from docx import Document
        from docx.shared import Pt
        import io

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        doc.add_heading(project["title"], level=0)
        if project.get("description"):
            doc.add_paragraph(project["description"])

        for sec in project.get("sections", []):
            doc.add_heading(sec["title"], level=1)
            content = sec.get("content", "") or "（待撰写）"
            for line in content.split("\n"):
                line = line.strip()
                if line.startswith("- "):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line:
                    doc.add_paragraph(line)

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    async def export_to_bibtex(self, project_id: str, user_id: str) -> str | None:
        """导出写作项目关联论文的 BibTeX。"""
        project = await self.get_project(project_id, user_id)
        if not project:
            return None

        from app.services.writing_service import WritingAssistantService

        papers = await self._collect_reference_papers(project)
        writer = WritingAssistantService(self.session)
        return "\n\n".join(writer._generate_bibtex(paper) for paper in papers)

    async def export_reference_list(self, project_id: str, user_id: str, style: str = "numeric") -> str | None:
        """导出可读参考文献列表。"""
        project = await self.get_project(project_id, user_id)
        if not project:
            return None
        papers = await self._collect_reference_papers(project)
        if not papers:
            return "暂无可导出的参考文献。请先在写作项目中关联证据论文或补充 References。"
        return self._format_reference_list(papers, style=style)

    async def build_export_readiness(self, project_id: str, user_id: str) -> dict | None:
        """生成投稿导出预检结果。"""
        project = await self.get_project(project_id, user_id)
        if not project:
            return None

        sections = project.get("sections") or []
        empty_sections = []
        short_sections = []
        citation_mentions = []
        unmatched_citations = []
        total_words = 0

        evidence = await self.get_evidence_cards(project_id, user_id)
        cards = (evidence or {}).get("cards") or []
        coverage = (evidence or {}).get("coverage") or {
            "total": 0, "local": 0, "external": 0, "bibtex_ready": 0,
        }
        metadata = project.get("metadata_json") or {}
        submission_profile = metadata.get("submission_profile") or {}

        for section in sections:
            title = section.get("title") or "Untitled"
            content = section.get("content") or ""
            word_count = self._count_words(content)
            total_words += word_count
            if not content.strip():
                empty_sections.append(title)
            elif word_count < 80 and title.lower() not in {"references", "related work comparison table"}:
                short_sections.append(title)
            for mention in self._extract_citation_mentions(content):
                citation_mentions.append({**mention, "section": title})
                if not self._resolve_citation_card(mention, cards):
                    unmatched_citations.append({"citation": mention["raw"], "section": title})

        papers = await self._collect_reference_papers(project)
        warnings = []
        if empty_sections:
            warnings.append(f"有 {len(empty_sections)} 个章节为空：{', '.join(empty_sections[:4])}")
        if short_sections:
            warnings.append(f"有 {len(short_sections)} 个章节内容偏短：{', '.join(short_sections[:4])}")
        if coverage.get("total", 0) == 0:
            warnings.append("项目没有证据卡，导出内容缺少可追溯论文支撑。")
        elif coverage.get("local", 0) < max(1, int(coverage.get("total", 0) * 0.5)):
            warnings.append("本地入库证据覆盖率偏低，建议先补全文/补向量后再定稿。")
        if coverage.get("total", 0) and coverage.get("bibtex_ready", 0) < coverage.get("total", 0):
            warnings.append("部分证据缺少可导出的 BibTeX，本次参考文献可能不完整。")
        if unmatched_citations:
            warnings.append(f"有 {len(unmatched_citations)} 个正文引用未匹配到证据卡。")
        if citation_mentions and not papers:
            warnings.append("正文已有引用标记，但没有解析到真实论文条目。")
        if not submission_profile:
            warnings.append("尚未绑定官方投稿模板；内置结构模板不能保证当前年度会议格式。")
        elif submission_profile.get("warnings"):
            warnings.append("投稿模板需要人工确认：" + " ".join(submission_profile.get("warnings", [])[:3]))

        status = "ready"
        if empty_sections:
            status = "incomplete"
        elif warnings:
            status = "needs_attention"

        return {
            "project_id": project_id,
            "status": status,
            "status_label": {
                "ready": "可导出",
                "needs_attention": "建议检查",
                "incomplete": "内容未完成",
            }[status],
            "warnings": warnings,
            "section_summary": {
                "total": len(sections),
                "empty": len(empty_sections),
                "short": len(short_sections),
                "total_words": total_words,
                "empty_sections": empty_sections,
                "short_sections": short_sections,
            },
            "evidence_coverage": coverage,
            "citation_summary": {
                "mentions": len(citation_mentions),
                "unmatched": len(unmatched_citations),
                "unmatched_items": unmatched_citations[:8],
            },
            "reference_summary": {
                "papers": len(papers),
                "bibtex_ready": coverage.get("bibtex_ready", 0),
            },
            "submission_profile": submission_profile or {
                "template_status": "missing",
                "status_label": "未绑定官方模板",
                "warnings": ["尚未上传或绑定官方投稿模板。"],
            },
        }

    async def build_publication_package(self, project_id: str, user_id: str) -> dict | None:
        """导出投稿包：正文、引用和导出预检信息。"""
        project = await self.get_project(project_id, user_id)
        if not project:
            return None
        markdown = await self.export_to_markdown(project_id, user_id) or ""
        latex = await self.export_to_latex(project_id, user_id) or ""
        bibtex = await self.export_to_bibtex(project_id, user_id) or ""
        references = await self.export_reference_list(project_id, user_id) or ""
        readiness = await self.build_export_readiness(project_id, user_id)
        base = self._safe_export_basename(project.get("title") or "writing-project")
        return {
            "project_id": project_id,
            "title": project.get("title") or "",
            "readiness": readiness,
            "formats": {
                "markdown": {"filename": f"{base}.md", "content": markdown},
                "latex": {"filename": f"{base}.tex", "content": latex},
                "bibtex": {"filename": f"{base}.bib", "content": bibtex},
                "references": {"filename": f"{base}_references.md", "content": references},
                "docx": {"filename": f"{base}.docx", "download_url": f"/api/writing/projects/{project_id}/export?format=docx"},
            },
        }

    async def build_workbench_summary(self, project_id: str, user_id: str) -> dict | None:
        """Build a project-first writing workbench summary with next actions."""
        project = await self.get_project(project_id, user_id)
        if not project:
            return None

        readiness = await self.build_export_readiness(project_id, user_id)
        evidence = await self.get_evidence_cards(project_id, user_id)
        progress = project.get("progress") or {}
        sections = project.get("sections") or []
        section_summary = (readiness or {}).get("section_summary") or {}
        evidence_coverage = (evidence or {}).get("coverage") or (readiness or {}).get("evidence_coverage") or {
            "total": 0,
            "local": 0,
            "external": 0,
            "bibtex_ready": 0,
        }
        citation_summary = (readiness or {}).get("citation_summary") or {}
        submission_profile = (readiness or {}).get("submission_profile") or {}
        warnings = list((readiness or {}).get("warnings") or [])
        stage = self._workbench_stage(progress, section_summary, evidence_coverage, readiness)
        next_actions = self._workbench_next_actions(
            project=project,
            stage=stage,
            readiness=readiness or {},
            evidence_coverage=evidence_coverage,
            citation_summary=citation_summary,
            submission_profile=submission_profile,
            sections=sections,
        )
        risk_level = self._workbench_risk_level(warnings, evidence_coverage, citation_summary, submission_profile)

        return {
            "project_id": project_id,
            "title": project.get("title") or "",
            "mode": "paper" if project.get("template_type") != "nsfc" else "grant",
            "stage": stage,
            "status": (readiness or {}).get("status") or "needs_attention",
            "status_label": (readiness or {}).get("status_label") or "建议检查",
            "risk_level": risk_level,
            "progress": {
                "percentage": progress.get("percentage") or 0,
                "completed_sections": progress.get("completed") or 0,
                "total_sections": progress.get("total") or len(sections),
                "total_words": progress.get("total_words") or section_summary.get("total_words") or 0,
                "empty_sections": section_summary.get("empty") or 0,
                "short_sections": section_summary.get("short") or 0,
            },
            "evidence": {
                **evidence_coverage,
                "status": (evidence or {}).get("evidence_status") or ("sufficient" if evidence_coverage.get("total") else "insufficient"),
            },
            "citations": {
                "mentions": citation_summary.get("mentions") or 0,
                "unmatched": citation_summary.get("unmatched") or 0,
            },
            "submission": {
                "venue": submission_profile.get("venue") or "",
                "year": submission_profile.get("year") or "",
                "template_status": submission_profile.get("template_status") or "missing",
                "status_label": submission_profile.get("status_label") or "未绑定官方模板",
                "template_source": submission_profile.get("template_source") or "",
                "warnings": submission_profile.get("warnings") or [],
            },
            "warnings": warnings[:8],
            "next_actions": next_actions,
            "quick_links": [
                {"key": "sections", "label": "继续编辑章节", "target": "sections"},
                {"key": "evidence", "label": "查看证据卡片", "target": "evidence"},
                {"key": "export", "label": "投稿导出预检", "target": "export"},
            ],
        }

    def _workbench_stage(
        self,
        progress: dict,
        section_summary: dict,
        evidence_coverage: dict,
        readiness: Optional[dict],
    ) -> dict:
        if not progress.get("total"):
            return {"key": "setup", "label": "搭建结构", "description": "先创建章节结构和写作目标。"}
        if section_summary.get("empty"):
            return {"key": "drafting", "label": "补齐初稿", "description": "仍有空章节，优先把论文结构写完整。"}
        if evidence_coverage.get("total", 0) == 0 or evidence_coverage.get("local", 0) < max(1, int(evidence_coverage.get("total", 0) * 0.5)):
            return {"key": "evidence", "label": "补强证据", "description": "证据覆盖不足，先补本地论文和证据对比表。"}
        if (readiness or {}).get("status") == "needs_attention":
            return {"key": "review", "label": "引用与格式检查", "description": "草稿已成形，进入引用、证据和模板预检。"}
        return {"key": "export", "label": "准备导出", "description": "基础条件较完整，可以生成导出包并人工复核。"}

    def _workbench_next_actions(
        self,
        *,
        project: dict,
        stage: dict,
        readiness: dict,
        evidence_coverage: dict,
        citation_summary: dict,
        submission_profile: dict,
        sections: list,
    ) -> list[dict]:
        actions: list[dict] = []
        empty_sections = (readiness.get("section_summary") or {}).get("empty_sections") or []
        short_sections = (readiness.get("section_summary") or {}).get("short_sections") or []
        if empty_sections:
            actions.append({
                "key": "fill-empty-section",
                "label": f"补齐空章节：{empty_sections[0]}",
                "priority": "high",
                "target": "sections",
                "reason": "项目还没有形成完整论文骨架。",
            })
        elif short_sections:
            actions.append({
                "key": "expand-short-section",
                "label": f"扩展偏短章节：{short_sections[0]}",
                "priority": "medium",
                "target": "sections",
                "reason": "部分章节还不足以支撑正式草稿。",
            })

        total = evidence_coverage.get("total") or 0
        local = evidence_coverage.get("local") or 0
        if total == 0:
            actions.append({
                "key": "add-evidence",
                "label": "从论文库或研究方向补证据",
                "priority": "high",
                "target": "evidence",
                "reason": "没有证据卡，Related Work 和引用校验无法闭环。",
            })
        elif local < max(1, int(total * 0.5)):
            actions.append({
                "key": "localize-evidence",
                "label": "把外部证据入库并补全文/向量",
                "priority": "high",
                "target": "evidence",
                "reason": "本地证据覆盖偏低，最终写作容易变成摘要推测。",
            })
        else:
            actions.append({
                "key": "build-evidence-table",
                "label": "生成 Related Work 证据对比表",
                "priority": "medium",
                "target": "evidence-table",
                "reason": "把证据角色、基线和支持关系整理进写作结构。",
            })

        if citation_summary.get("unmatched", 0) > 0:
            actions.append({
                "key": "fix-unmatched-citations",
                "label": "修复未匹配引用",
                "priority": "high",
                "target": "citations",
                "reason": "正文引用没有对应证据卡，导出前必须确认。",
            })
        elif sections:
            actions.append({
                "key": "check-section-citations",
                "label": "逐章节运行引用校验",
                "priority": "medium",
                "target": "citations",
                "reason": "检查引用是否真实存在并能支撑句子。",
            })

        if not submission_profile or submission_profile.get("template_status") == "missing":
            actions.append({
                "key": "bind-submission-template",
                "label": "上传并绑定会议官方模板",
                "priority": "medium",
                "target": "submission-template",
                "reason": "内置结构模板不能保证当前年度投稿格式。",
            })

        if not actions:
            actions.append({
                "key": "build-export-package",
                "label": "生成投稿导出包",
                "priority": "medium",
                "target": "export",
                "reason": "草稿基础状态良好，可以导出 Markdown/BibTeX/Word 后人工复核。",
            })
        return actions[:5]

    def _workbench_risk_level(
        self,
        warnings: list[str],
        evidence_coverage: dict,
        citation_summary: dict,
        submission_profile: dict,
    ) -> str:
        if evidence_coverage.get("total", 0) == 0 or citation_summary.get("unmatched", 0) > 0:
            return "high"
        if len(warnings) >= 3 or not submission_profile or submission_profile.get("template_status") == "missing":
            return "medium"
        if warnings:
            return "low"
        return "clear"

    # --- 进度统计 ---

    def _calc_progress(self, sections: list) -> dict:
        """计算项目进度统计。"""
        total = len(sections)
        if total == 0:
            return {"percentage": 100, "completed": 0, "total": 0}

        completed = sum(1 for s in sections if s.status == "complete")
        return {
            "percentage": round(completed / total * 100),
            "completed": completed,
            "total": total,
            "total_words": sum(s.word_count or 0 for s in sections),
        }

    async def _collect_reference_papers(self, project: dict) -> list:
        """从项目元数据和 References 章节解析真实论文。"""
        from app.db.models.paper import Paper

        metadata = project.get("metadata_json") or {}
        paper_ids = list(dict.fromkeys(str(item) for item in (metadata.get("recommended_paper_ids") or []) if item))
        arxiv_ids = list(dict.fromkeys(str(item) for item in (metadata.get("recommended_arxiv_ids") or []) if item))
        dois = []

        references = "\n".join(
            section.get("content") or ""
            for section in project.get("sections", [])
            if section.get("title", "").lower() == "references"
        )
        for match in re.findall(r"Paper ID:\s*([0-9a-fA-F-]{32,36})", references):
            if match not in paper_ids:
                paper_ids.append(match)
        for match in re.findall(r"arXiv[:\s]+([0-9]{4}\.[0-9]{4,5}(?:v\d+)?)", references, flags=re.I):
            if match not in arxiv_ids:
                arxiv_ids.append(match)
        for match in re.findall(r"DOI[:\s]+([^\s,;]+)", references, flags=re.I):
            doi = match.strip().rstrip(".")
            if doi and doi not in dois:
                dois.append(doi)

        papers = []
        seen = set()
        for pid in paper_ids:
            try:
                result = await self.session.execute(select(Paper).where(Paper.id == UUID(pid)))
                paper = result.scalar_one_or_none()
                if paper and str(paper.id) not in seen:
                    papers.append(paper)
                    seen.add(str(paper.id))
            except ValueError:
                continue

        for arxiv_id in arxiv_ids:
            result = await self.session.execute(select(Paper).where(Paper.arxiv_id == arxiv_id))
            paper = result.scalar_one_or_none()
            if paper and str(paper.id) not in seen:
                papers.append(paper)
                seen.add(str(paper.id))

        for doi in dois:
            result = await self.session.execute(select(Paper).where(Paper.doi == doi))
            paper = result.scalar_one_or_none()
            if paper and str(paper.id) not in seen:
                papers.append(paper)
                seen.add(str(paper.id))

        return papers

    def _format_reference_list(self, papers: list, style: str = "numeric") -> str:
        rows = []
        for index, paper in enumerate(papers, 1):
            prefix = f"[{index}] " if style == "numeric" else ""
            authors = self._format_authors(getattr(paper, "authors", None))
            year = getattr(paper, "year", None) or "n.d."
            ids = []
            if getattr(paper, "arxiv_id", None):
                ids.append(f"arXiv:{paper.arxiv_id}")
            if getattr(paper, "doi", None):
                ids.append(f"DOI:{paper.doi}")
            suffix = f" {'; '.join(ids)}." if ids else ""
            rows.append(f"{prefix}{authors} ({year}). {paper.title}.{suffix}")
        return "\n\n".join(rows)

    def _format_authors(self, authors) -> str:
        if isinstance(authors, list):
            values = [str(item) for item in authors if item]
        elif isinstance(authors, dict):
            raw = authors.get("names") or authors.get("authors") or authors.get("list") or []
            values = [str(item) for item in raw] if isinstance(raw, list) else [str(raw)]
        elif isinstance(authors, str):
            values = [authors]
        else:
            values = []
        return ", ".join(values[:6]) or "Unknown Authors"

    def _count_words(self, text: str) -> int:
        ascii_words = re.findall(r"[A-Za-z0-9_]+", text or "")
        cjk_chars = re.findall(r"[\u4e00-\u9fff]", text or "")
        return len(ascii_words) + len(cjk_chars)

    def _safe_export_basename(self, title: str) -> str:
        safe = re.sub(r"[^A-Za-z0-9._-]+", "_", title.strip()).strip("_")
        return (safe or "writing_project")[:80]

    # --- Helper ---

    def _project_to_dict(self, project) -> dict:
        sections = sorted(
            [self._section_to_dict(s) for s in (project.sections or [])],
            key=lambda s: s["order"],
        )
        return {
            "id": str(project.id),
            "title": project.title,
            "description": project.description,
            "template_type": project.template_type,
            "status": project.status,
            "metadata_json": project.metadata_json or {},
            "sections": sections,
            "progress": self._calc_progress(project.sections or []),
            "created_at": project.created_at.isoformat() if project.created_at else "",
            "updated_at": project.updated_at.isoformat() if project.updated_at else "",
        }

    def _section_to_dict(self, section) -> dict:
        return {
            "id": str(section.id),
            "title": section.title,
            "content": section.content,
            "order": section.order,
            "status": section.status,
            "word_count": section.word_count or 0,
        }

    def _evidence_item_to_card(self, item: dict, paper, index: int) -> dict:
        local_id = item.get("imported_paper_id") or item.get("paper_id")
        has_local_id = bool(local_id and self._looks_like_uuid(str(local_id)))
        if paper:
            local_status = "local"
            local_status_label = "已入库"
        elif has_local_id:
            local_status = "unknown"
            local_status_label = "本地记录缺失"
        else:
            local_status = "external"
            local_status_label = "未入库"
        category = item.get("category") or item.get("role") or "supporting_evidence"
        role, role_label = self._normalize_evidence_role(category)
        title = item.get("title") or getattr(paper, "title", None) or "Untitled evidence"
        return {
            "id": f"evidence:{item.get('paper_id') or item.get('arxiv_id') or index}",
            "index": index,
            "citation_marker": f"[{index}]",
            "title": title,
            "year": item.get("year") or getattr(paper, "year", None),
            "authors": self._authors_to_text(getattr(paper, "authors", None) or item.get("authors")),
            "source": item.get("source") or getattr(paper, "source", None) or "evidence",
            "source_label": item.get("source") or ("本地论文库" if local_status == "local" else "外部证据"),
            "source_kind": item.get("category") or "evidence",
            "paper_id": str(getattr(paper, "id", local_id)) if (local_status in {"local", "unknown"} and local_id) else None,
            "external_paper_id": item.get("paper_id") if not self._looks_like_uuid(str(item.get("paper_id") or "")) else None,
            "arxiv_id": item.get("arxiv_id") or getattr(paper, "arxiv_id", None),
            "doi": item.get("doi") or getattr(paper, "doi", None),
            "role": role,
            "role_label": role_label,
            "snippet": item.get("relevance") or item.get("abstract_excerpt") or (getattr(paper, "abstract", "") or "")[:240],
            "local_status": local_status,
            "local_status_label": local_status_label,
            "bibtex_ready": local_status == "local",
        }

    def _paper_to_evidence_card(self, paper, index: int) -> dict:
        return {
            "id": f"paper:{paper.id}",
            "index": index,
            "citation_marker": f"[{index}]",
            "title": paper.title,
            "year": paper.year,
            "authors": self._authors_to_text(paper.authors),
            "source": paper.source,
            "source_label": "本地论文库",
            "source_kind": "recommended_paper",
            "paper_id": str(paper.id),
            "external_paper_id": None,
            "arxiv_id": paper.arxiv_id,
            "doi": paper.doi,
            "role": "supporting_evidence",
            "role_label": "支持证据",
            "snippet": (paper.abstract or "")[:240] or "该论文暂无摘要，建议补摘要/补全文后再做精细引用校验。",
            "local_status": "local",
            "local_status_label": "已入库",
            "bibtex_ready": True,
        }

    def _normalize_evidence_role(self, role: str) -> tuple[str, str]:
        mapping = {
            "seed": ("supporting_evidence", "核心证据"),
            "inspiration": ("supporting_evidence", "灵感证据"),
            "background": ("background", "背景资料"),
            "baseline": ("baseline_method", "基线方法"),
            "baseline_method": ("baseline_method", "基线方法"),
            "counterexample": ("counterexample", "反例/局限"),
            "limitation": ("counterexample", "反例/局限"),
        }
        return mapping.get((role or "").lower(), ("supporting_evidence", "支持证据"))

    def _authors_to_text(self, authors) -> str:
        if isinstance(authors, list):
            return ", ".join(str(author) for author in authors[:4])
        if isinstance(authors, dict):
            values = authors.get("names") or authors.get("authors") or []
            if isinstance(values, list):
                return ", ".join(str(author) for author in values[:4])
        return str(authors or "")

    def _evidence_card_key(self, card: dict) -> str:
        return (
            card.get("paper_id")
            or self._normalize_arxiv_id(card.get("arxiv_id") or "")
            or (card.get("doi") or "").lower()
            or (card.get("title") or "").lower()
        )

    def _normalize_arxiv_id(self, value: str) -> str:
        return re.sub(r"^arxiv[:\s]*", "", value or "", flags=re.I).strip().lower()

    def _extract_citation_mentions(self, text: str) -> list[dict]:
        mentions: list[dict[str, Any]] = []
        for match in re.finditer(r"\[(\d+)\]", text or ""):
            mentions.append({
                "type": "index",
                "raw": match.group(0),
                "value": int(match.group(1)),
                "start": match.start(),
                "end": match.end(),
            })
        for match in re.finditer(r"arXiv[:\s]+([0-9]{4}\.[0-9]{4,5}(?:v\d+)?)", text or "", flags=re.I):
            mentions.append({
                "type": "arxiv",
                "raw": match.group(0),
                "value": match.group(1),
                "start": match.start(),
                "end": match.end(),
            })
        for match in re.finditer(r"Paper ID:\s*([0-9a-fA-F-]{32,36})", text or "", flags=re.I):
            mentions.append({
                "type": "paper_id",
                "raw": match.group(0),
                "value": match.group(1),
                "start": match.start(),
                "end": match.end(),
            })
        return sorted(mentions, key=lambda item: item["start"])

    def _resolve_citation_card(self, mention: dict, cards: list[dict]) -> dict | None:
        if mention["type"] == "index":
            index = mention["value"]
            return cards[index - 1] if 1 <= index <= len(cards) else None
        if mention["type"] == "arxiv":
            normalized = self._normalize_arxiv_id(mention["value"])
            return next((card for card in cards if self._normalize_arxiv_id(card.get("arxiv_id") or "") == normalized), None)
        if mention["type"] == "paper_id":
            value = str(mention["value"]).lower()
            return next((card for card in cards if str(card.get("paper_id") or "").lower() == value), None)
        return None

    def _citation_context(self, text: str, start: int, end: int) -> str:
        before = max(text.rfind("。", 0, start), text.rfind(".", 0, start), text.rfind("\n", 0, start))
        after_candidates = [pos for pos in [text.find("。", end), text.find(".", end), text.find("\n", end)] if pos != -1]
        after = min(after_candidates) if after_candidates else min(len(text), end + 240)
        return text[before + 1:after + 1].strip()[:500] or text.strip()[:500]

    def _citation_summary(self, checks: list[dict]) -> dict:
        counts = {"strong": 0, "partial": 0, "weak": 0, "missing": 0, "unchecked": 0}
        for check in checks:
            status = check.get("status")
            counts[status if status in counts else "weak"] += 1
        total = len(checks)
        verified = counts["strong"] + counts["partial"]
        return {
            "total": total,
            **counts,
            "citation_coverage": round(verified / total, 3) if total else 0,
            "evidence_warning": counts["weak"] > 0 or counts["missing"] > 0 or counts["unchecked"] > 0,
        }

    def _recommend_evidence_cards(self, cards: list[dict], limit: int = 3) -> list[dict]:
        return [
            {
                "citation_marker": card.get("citation_marker"),
                "title": card.get("title"),
                "role_label": card.get("role_label"),
                "local_status_label": card.get("local_status_label"),
            }
            for card in cards[:limit]
        ]

    def _evidence_writing_use(self, card: dict) -> str:
        status = card.get("local_status")
        role = card.get("role_label") or "支持证据"
        if status == "local":
            return f"可作为{role}写入 Related Work；定稿前建议继续运行章节引用校验。"
        if status == "unknown":
            return "记录过本地 ID 但当前论文缺失，建议重新入库或替换引用。"
        return f"可作为{role}的线索；尚未入库，不能当作已验证证据。"
